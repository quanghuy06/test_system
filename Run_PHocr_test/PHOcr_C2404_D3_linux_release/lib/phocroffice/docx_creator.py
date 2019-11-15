# -*- encoding: utf-8 -*-
# Copyright (C) 2017 TSDV TTEC.  All rights reserved.

from __future__ import division

import warnings

import time
import os

warnings.filterwarnings('ignore', category=DeprecationWarning)

import sys

from docx import Document
from docx.enum.section import \
    WD_ORIENTATION as WD_ORIENT, \
    WD_TEXT_DIRECTION as WD_TEXT_DIRECTION
from phocr_shared.shared import Pt, Inches, Pixel, Emu
from docx.box import ratio_overlap_photo

from docx.common import Common
from phocr_elements.office_object_tag import OfficeObjectTag
from phocr_shared.object_info import ObjectInfo
from phocr_shared.photo_attribute import PhotoAttribute
from phocr_shared.table_attribute import TableAttribute
from phocr_shared.textbox_attribute import TextBoxAttribute
from phocr_shared.table_cell_info import TableCellInfo
from phocr_shared.font_info import FontInfo
from phocr_shared.word_property import WordProperty
from phocr_shared.word_format import WordFormat
from docx.table_processing import TableProcessing
from phocr_shared.phocr_config import PHOcrConstant
from phocr_shared.phocr_common import PHOcrCommon, TextDirection, DocumentType
from phocr_shared.bullets_and_numberings import update_bullets_and_numbering, \
    BulletNumberingContainer
from phocr_shared.alignment import string_to_alignment_docx, analysis_indent, \
    indent_one_line_paragraph, \
    calculate_left_indent
from phocr_shared.text_direction import OrientationTransformer
from phocr_shared import text_box
from phocr_shared.logger import OCRLogLevel
from phocr_shared.element_helper import paragraph_to_string
from phocr_shared.box_container import BoxContainer
from phocr_shared.ocr_box import OCRBox
from base_office_creator import BaseOfficeCreator
from phocr_shared.phocr_dpi import PHOcrDPI
from phocr_shared.text_direction import TextlineOrder

STANDARD_DEVIATION_THRESHOLD_SEPARATED = 10
STANDARD_DEVIATION_THRESHOLD_TO_SEPARATED = 30
THRESHOLD_ALIGNMENT_NEAR_IMAGE = 8
THRESHOLD_ALIGNMENT_STD = 5
THRESHOLD_UP_XHEIGHT = 3 / 4
THRESHOLD_DOWN_XHEIGHT = 2 / 5
BLINKING_CURSOR_HEIGHT = 12


class DocxCreator(BaseOfficeCreator):
    """Creator to generate .docx file"""

    def __init__(self, filename, language, debug, working_directory):
        super(DocxCreator, self).__init__(language, working_directory, debug)
        self.all_box_photo = []
        self.width = 0
        self.height = 0
        self.x_height = 1
        self.baseline = 1
        self.dpi = PHOcrDPI(PHOcrConstant.DEFAULT_DPI, PHOcrConstant.DEFAULT_DPI)  # Set default dpi
        self.area_orders = []
        self.is_complicated = False
        self.is_east_asia = False
        # Set the value for self.is_east_asia in case we are working
        # with East Asian language such as Chinese, Japanese or Korean
        if Common.is_cjk_language(self.language):
            self.is_east_asia = True
        self.document_type = DocumentType.DOCX
        self.output_file = filename + '.docx'

        # Margin
        self.margin_top = 0
        self.margin_bottom = 0
        self.margin_left = 0
        self.margin_right = 0

        self.bullet_numberings = BulletNumberingContainer()
        self.is_list_para = False

        # Direction
        self.transformer = OrientationTransformer()

        # Default text line direction
        self.text_direction = TextDirection.LRTB

        # Standard paper width (Inch unit)
        self.paper_width = 0

        # Standard paper height (Inch unit)
        self.paper_height = 0

        # Standard paper code, used to change the MS Office XML structure
        self.paper_code = 0

        # Standard name of this paper size (A4, A3, Letter...)
        self.paper_name = ""

        # Paper orientation
        self.paper_orientation = ""

        # Flag to check if page is standard paper size
        self.is_using_standard_size = True

    def set_word_style(self, p, run, word):
        font = run.font
        if word.color is not None:
            font.color.rgb = PHOcrCommon.get_color(word.color)
        if word.highlight_color is not None:
            current_highlight = Common.get_highlight_color(word.highlight_color)
            font.highlight_color = current_highlight
            if len(p.runs) > 2:
                r_2_highlight = p.runs[-3].font.highlight_color
                if r_2_highlight == current_highlight:
                    p.runs[-2].font.highlight_color = current_highlight
        if word.font and len(word.font) > 0:
            font.name = word.font
            if self.is_east_asia:
                font.east_asia = word.font
        if word.size and len(word.size) > 0:
            font.size = Pt(float(word.size))
        if word.bold:
            run.bold = True
        if word.italic:
            run.italic = True
        if word.underline:
            run.underline = True

    def _set_space_before(self, p, word, previous_word, word_count,
                          is_same_format=None):
        """
        Set space before properties for current word based on space_before tag

        Parameters
        ----------
        p
            Paragraph element from python docx library
        word
            Current word element from element tree
        previous_word
            Previous word element from element tree
        word_count
            Position of this word in line, start count from 0
        is_same_format
            The word and previous word are same format.
            With the condition, the space is belong to the word.
        Returns
        -------

        """
        # Instance two Word object from element tree object for easy processing
        word_object = WordProperty(word)
        previous_word_object = WordProperty(previous_word)

        # Get the attribute of current word
        space_before = word_object.spaces_before


        # This is the full-width space mark used in japanese language
        space_mark = '　'

        if len(p.runs) > 0 and space_before > 0 and is_same_format:
            run = p.runs[-1]
            run.text = run.text + space_mark * space_before
        else:
            if space_before > 0:
                run_space = p.add_run()
                run_space.text = space_mark * space_before
                # If current word is the first word, then we will set style for
                # this space based on first word style except underline attribute
                if word_count == 0:
                    space_font_info = FontInfo(word)
                    space_font_info.underline = 'False'
                    space_color = word_object.color
                else:
                    # Check style of this word and the previous word
                    # If they are all underline then space between is
                    # set underline too. Else, space style is
                    # set to previous word style except underline
                    space_font_info = FontInfo(previous_word)
                    if space_font_info.underline == 'True' and \
                            word_object.underline == 'True':
                        pass
                    else:
                        space_font_info.underline = 'False'
                    space_color = previous_word_object.color
                Common.set_text_style_for_run(run_space,
                                              space_font_info,
                                              space_color)

    def generate_word(
            self,
            p,
            par,
            word,
            prev_word,
            word_count,
            length=1,
            index=1,
            next_word=None,
            box_container=None):
        """
        Generate a word into MS Word Document representation

        Parameters
        ----------
        p
        par
        word
        prev_word
        word_count
        length
        index
        next_word
        box_container

        Returns
        -------

        """
        word_object = WordFormat(word)
        previous_word_object = WordFormat(prev_word)
        same_format_preword = word_object.is_same_format(previous_word_object)

        # When export one word to MS Word, we must to export the space before
        # This space before is set from space_before tag in each word element
        # and space's style is set from previous word and current word
        if self.is_east_asia:
            self._set_space_before(p, word, prev_word, word_count,
                                   same_format_preword)

        if len(p.runs) > 0 and same_format_preword:
            run = p.runs[-1]
            run.text = run.text + word.value
        else:
            run = p.add_run()
            self.set_word_style(p, run, word)
            run.add_text(word.value)

        next_word_object = WordFormat(next_word)
        same_format_next_word = word_object.is_same_format(next_word_object)
        # If working with languages that are not East Asian language, we'll use
        # the old algorithm. It sets space after each word.
        if not self.is_east_asia:
            if index < length:
                Common.set_space_after(
                    p,
                    par,
                    word,
                    self.dpi,
                    self.transformer,
                    next_word,
                    box_container,
                    same_format_next_word)
        return index

    def generate_line(
            self,
            p,
            par,
            line,
            is_list_para,
            is_first_line,
            box_container=None,
            end_word_pre_line=None):
        """
        Generate a line into MS Word Document representation

        Parameters
        ----------
        p
        par
        line
        is_list_para
        is_first_line
        box_container
        end_word_pre_line
            It is a word, the word is end word, belong to previous line
        Returns
        -------

        """
        # Reset word_count wehn we start generate a new line
        index = 1
        word_count = 0
        for word in line.words:
            prev_word = None
            if 0 < word_count < len(line.words):
                prev_word = line.words[word_count - 1]
            if word_count == 0:
                prev_word = end_word_pre_line
            next_word = None
            if index < len(line.words):
                next_word = line.words[index]
            ignore = False
            if is_first_line and is_list_para and index == 1:
                ignore = True
            if not ignore:
                index = self.generate_word(
                    p, par, word, prev_word,
                    word_count, len(line.words),
                    index, next_word,
                    box_container=box_container
                )
            word_count += 1
            index += 1

    def generate_paragraph(
            self,
            p,
            par,
            is_list_para,
            box_container=None,
            block=None):
        """
        Generate a paragraph into MS Word Document representation

        Parameters
        ----------
        p
        par
        is_list_para
        box_container
        block

        Returns
        -------

        """
        if par.background_color is not None:
            p.paragraph_format.shd.rgb = PHOcrCommon.get_color(
                par.background_color
            )
        # Set alignment for paragraph
        p.alignment = string_to_alignment_docx(par.alignment)

        i = 0
        for line in par.lines:
            end_word_pre_line = None
            if (i - 1 >= 0) and (i - 1 < len(par.lines)):
                end_word_pre_line = par.lines[i - 1].words[-1]

            self.generate_line(
                p,
                par,
                line,
                is_list_para,
                i == 0,
                box_container,
                end_word_pre_line)

            last_word = line.words[-1]
            if i < len(par.lines) - 1:
                first_word_next_line = par.lines[i + 1].words[0]
                last_word_object = WordFormat(last_word)
                first_word_next_line_object = WordFormat(first_word_next_line)
                is_same_format = last_word_object.is_same_format(
                    first_word_next_line_object)
                if not self.is_east_asia:
                    Common.set_space_after(
                        p,
                        par,
                        last_word,
                        self.dpi,
                        self.transformer,
                        first_word_next_line,
                        box_container,
                        is_same_format)
            i += 1

        self.setting_line_spacing(p, par)

        # Calculate indentation
        (left_indent, right_indent, first_line_indent) = \
            analysis_indent(p, par, self.dpi, Pixel,
                            block=block,
                            page={
                                'margin_left':
                                    Pixel(box_container.left_margin(),
                                          self.dpi.horizontal_resolution).emu,
                                'margin_right':
                                    Pixel(box_container.right_margin(),
                                          self.dpi.horizontal_resolution).emu,
                                'margin_top':
                                    Pixel(box_container.top_margin(),
                                          self.dpi.vertical_resolution).emu,
                                'margin_bottom':
                                    Pixel(box_container.bottom_margin(),
                                          self.dpi.vertical_resolution).emu,
                                'width':
                                    Pixel(box_container.parent_w_px,
                                          self.dpi.horizontal_resolution).emu,
                                'height':
                                    Pixel(box_container.parent_h_px,
                                          self.dpi.vertical_resolution).emu,
                                'text_line_order':
                                    self.transformer.text_line_order,
                                'writing_direction':
                                    self.transformer.writing_direction,
                                'orientation':
                                    self.transformer.orientation
                            })

        # Correct left, right indent with one line paragraph
        left_indent, right_indent = \
            indent_one_line_paragraph(par, left_indent,
                                      right_indent,
                                      self.transformer,
                                      Pixel)

        p.paragraph_format.first_line_indent = first_line_indent
        p.paragraph_format.left_indent = \
            calculate_left_indent(p.paragraph_format, left_indent)
        p.paragraph_format.right_indent = right_indent

        # Logging
        self.file_logger.log(
            OCRLogLevel.LOG_LEVEL_DEBUG,
            'Paragraph: {0}\n'
            '   left_indent_before: {left_indent_before}\n'
            '   left_indent_after: {left_indent_after}\n'
            '   right_indent_before: {right_indent_before}\n'
            '   right_indent_after: {right_indent_after}\n'
            '   first_line_indent_before: {first_line_indent_before}\n'
            '   first_line_indent_after: {first_line_indent_after}\n'
            '   page: {page}\n'
            .format(
                paragraph_to_string(par),
                left_indent_before=par.left_indent,
                left_indent_after=Emu(
                    p.paragraph_format.left_indent
                ).inches * self.dpi.horizontal_resolution,
                right_indent_before=par.right_indent,
                right_indent_after=Emu(
                    p.paragraph_format.right_indent
                ).inches * self.dpi.horizontal_resolution,
                first_line_indent_before=par.first_line_indent,
                first_line_indent_after=Emu(
                    p.paragraph_format.first_line_indent
                ).inches * self.dpi.horizontal_resolution,
                page={
                    'margin_left': box_container.left_margin(),
                    'margin_right': box_container.right_margin(),
                    'margin_top': box_container.top_margin(),
                    'margin_bottom': box_container.bottom_margin(),
                    'width': box_container.parent_w_px,
                    'height': box_container.parent_h_px,
                    'text_line_order': self.transformer.text_line_order,
                    'writing_direction': self.transformer.writing_direction,
                    'orientation': self.transformer.orientation
                }
            )
        )

    def setting_line_spacing(self, p, par):
        if len(par.lines) > 1:
            line_spacing_pix = PHOcrCommon.calculate_line_spacing(
                par,
                self.text_direction,
                self.document_type
            )

            # Initialize default value of the right resolution
            resolution = PHOcrConstant.DEFAULT_DPI
            if self.text_direction == TextDirection.LRTB:
                resolution = self.dpi.horizontal_resolution
            elif self.text_direction == TextDirection.TBRL:
                resolution = self.dpi.vertical_resolution
            p.paragraph_format.line_spacing = Pixel(line_spacing_pix, resolution)
        else:
            # Default is not 1.0 because we have some incline paragraph
            # 0.1 for more suitable with word format, because bounding box of
            # OCR output are strictly fit with character
            # --> 0.1 will take care the error from OCR (eg. binarization, ... )
            p.paragraph_format.line_spacing = 1.1

    def _get_table_infor(self, document, page, element, table_row, table_col):
        table_info = []

        for _ in range(table_row):
            row_info = []
            for _ in range(table_col):
                cell_info = TableCellInfo('', '', '')
                row_info.append(cell_info)
            table_info.append(row_info)

        for row in element.rows:
            for column in row.cells:
                self.generate_content_for_cell(
                    document,
                    page,
                    column,
                    row.id - 1,
                    table_info
                )
        return table_info

    def create_table_object(self, document, page, element, column_w, column_h):
        table_pos = PHOcrCommon.get_object_position(element)
        table_info = self._get_table_infor(document, page, element,
                                           element.num_rows, element.num_cells)

        # Build table grid
        tbl = TableAttribute(
            self.dpi,
            element.num_rows,
            element.num_cells,
            table_pos.x,
            table_pos.y,
            table_pos.w,
            table_pos.h,
            column_w,
            column_h,
            table_info
        )
        table = document.add_updated_table_with_table_paras(tbl=tbl)
        table.style = 'Table Grid'
        table.cell_margin(0, 0, 0, 0)
        return table

    def generate_table(
            self, document, page, element,
            column_w, column_h):
        """
        Generate a table into MS Word Document representation

        Parameters
        ----------
        document
        page
        element
        column_w
        column_h

        Returns
        -------

        """
        time_start = time.time()
        table = self.create_table_object(document, page,
                                         element, column_w, column_h)

        if self.debug:
            t2 = time.time()
            print("Total create_table_object %.2f seconds" % (t2 - time_start))

        # NamLD: Refer: https://github.com/python-openxml/python-docx/issues/174
        # Accessing by using table.cell() will take a lot of time. Therefore,
        # we need access to table._cells directly
        table_cells = table._cells
        num_columns = element.num_cells
        for row in element.rows:
            for column in row.cells:
                row_start = row.id - 1
                col_start = column.id - 1
                if column.col_span > 1 or column.row_span > 1:
                    cell_start = table_cells[row_start*num_columns + col_start]
                    cell_end = table_cells[
                        (row_start + column.row_span - 1)*num_columns +
                        (col_start + column.col_span - 1)
                    ]
                    cell_start.merge(cell_end)

        if self.debug:
            t2 = time.time()
            print("Total merge %.2f seconds" % (t2 - time_start))

    def generate_element_in_cell(self, document, page,
                                 table_info, column_xml,
                                 row_id, col_id, row_span, col_span):
        # Assign other properties of the cell
        for carea in column_xml.careas:
            c_area = self.generate_text_for_cell(
                document,
                page,
                column_xml,
                carea
            )
            table_info[row_id][col_id].add_carea_for_cell(c_area)
        border_info = TableProcessing.generate_border_info_for_cell(
            column_xml
        )
        for col_idx in range(col_span):
            table_info[row_id][col_id + col_idx].add_border_for_cell(
                border_info
            )
        for row_idx in range(row_span):
            table_info[row_id + row_idx][col_id].add_border_for_cell(
                border_info
            )

    def generate_content_for_cell(
            self,
            document,
            page,
            column_xml,
            row_id,
            table_info):
        col_id = column_xml.id - 1
        # Assign the background color for the cell
        TableProcessing.assign_cell_color(
            table_info[row_id][col_id],
            column_xml
        )
        self.generate_element_in_cell(document, page,
                                      table_info, column_xml,
                                      row_id, col_id,
                                      column_xml.row_span, column_xml.col_span)

    def generate_text_for_cell(self, document, page, column_xml, carea_xml):
        cell_box = OCRBox.from_xml_element(column_xml, self.dpi, Pixel)
        page_box = OCRBox.from_xml_element(page, self.dpi, Pixel)
        box_container = BoxContainer(
            box=cell_box,
            parent_box=page_box
        )
        paragraphs, _ = self.generate_carea(
            document=document,
            page=column_xml,
            block=carea_xml,
            first_paragraph_available=False,
            box_container=box_container
        )
        return paragraphs

    def generate_photo(self, document, photo, wrap):
        """
        Generate a photo into MS Word Document representation

        Parameters
        ----------
        document
        photo
        wrap

        Returns
        -------

        """
        # Only add photo if it is bigger than threshold
        if PHOcrCommon.is_photo_ok(photo.h, photo.w, self.x_height):
            p = Common.get_anchor_paragraph(document)
            r = p.add_run()
            absolute_photo_path = os.path.join(self.working_directory, photo.path)
            photo = PhotoAttribute(
                self.dpi,
                photo.x,
                photo.y,
                photo.w,
                photo.h,
                absolute_photo_path,
                wrap,
                photo.inside_table
            )
            r.add_photo(photo)

    def get_next_carea(self, page, block):
        for i in range(len(self.area_orders)):
            area_obj = self.area_orders[i]
            if block.id == area_obj.obj_id and i < len(self.area_orders) - 1:
                next_area_obj = self.area_orders[i + 1]
                for element in page.elements:
                    if element.tag == OfficeObjectTag.CAREA and \
                            element.id == next_area_obj.obj_id and \
                            element.tag == next_area_obj.tag:
                        if element.is_noise or \
                                ratio_overlap_photo(element,
                                                    self.all_box_photo) >= 1:
                            return self.get_next_carea(page, element)
                        else:
                            return element
                    elif element.tag == OfficeObjectTag.TABLE and \
                            element.id == next_area_obj.obj_id and \
                            element.tag == next_area_obj.tag:
                        return element

    def get_saved_space_after(
            self,
            page,
            block,
            box_container,
            par,
            paragraph_orders,
            i):
        saved_space_after = 0
        # If current paragraph is not last paragraph in carea
        # We'll calculate space before with the next paragraph
        if i < len(paragraph_orders) - 1:
            next_para_obj = paragraph_orders[i + 1]
            for next_para in block.paragraphs:
                if next_para.id == next_para_obj.obj_id:
                    next_par = next_para
                    saved_space_after = Common.calculate_space_after(
                        par,
                        next_par,
                        self.text_direction
                    )
                    break
        elif box_container.is_whole_page_box:
            # If current paragraph is last paragraph in carea and
            # not inside a textbox
            # We'll calculate space before for the first paragraph
            # of next carea or table with condition is
            # current careais not last carea in page
            element = self.get_next_carea(page, block)
            if element is not None:
                if element.tag == OfficeObjectTag.CAREA:
                    next_carea = element
                    # Firstly, store position of each paragraph of
                    # content area
                    next_paragraph_orders = []
                    for next_para in next_carea.paragraphs:
                        para_info = ObjectInfo(
                            next_para.id,
                            next_para.x,
                            next_para.y,
                            next_para.tag
                        )
                        next_paragraph_orders.append(para_info)
                    if self.text_direction == TextDirection.TBRL:
                        # Next, sort by x and reverse (DEC)
                        next_paragraph_orders = sorted(
                            next_paragraph_orders,
                            key=lambda p_element: p_element.x,
                            reverse=True
                        )
                    next_para_obj = next_paragraph_orders[0]
                    for next_para in next_carea.paragraphs:
                        if next_para.id == next_para_obj.obj_id:
                            saved_space_after = Common.calculate_space_after(
                                par,
                                next_para,
                                self.text_direction
                            )
                            break
                elif element.tag == OfficeObjectTag.TABLE:
                    saved_space_after = Common.calculate_space_after(
                        par,
                        element,
                        self.text_direction
                    )
        return saved_space_after

    def _set_space_after(
            self,
            page,
            block,
            box_container,
            p,
            par,
            paragraph_orders,
            i):
        # If current paragraph is not last paragraph in carea
        # We'll calculate space before with the next paragraph
        saved_space_after = self.get_saved_space_after(
            page,
            block,
            box_container,
            par,
            paragraph_orders,
            i
        )

        # Initialize default value of the right resolution
        resolution = PHOcrConstant.DEFAULT_DPI
        if self.text_direction == TextDirection.LRTB:
            resolution = self.dpi.vertical_resolution
        elif self.text_direction == TextDirection.TBRL:
            resolution = self.dpi.horizontal_resolution
        tolerance = Inches(2 * BLINKING_CURSOR_HEIGHT / resolution)
        real_space_after = ((Inches(saved_space_after / resolution)
                             - tolerance))
        if real_space_after < 0:
            real_space_after = 0
        p.paragraph_format.space_after = real_space_after

    def generate_paragraph_inside_carea(
            self,
            document,
            page,
            block,
            first_paragraph_available,
            box_container,
            par,
            text_angle_histogram,
            paragraph_orders,
            i):
        carea_x = block.x
        overlap_v, near_image, region_height = \
            Common.calculate_para_infor(page, par, self.transformer)
        is_near_image = Common.get_near_image(
            par,
            self.transformer,
            overlap_v,
            near_image,
            region_height
        )

        if first_paragraph_available:
            space_after = \
                Common.calculate_space_after_first_paragraph(
                    page,
                    par,
                    self.margin_top
                )
            first_paragraph = document.paragraphs[-1]
            if space_after > 10:
                # If there are space after need for first paragraph
                # The next paragraph will need to be new
                p = document.add_paragraph()
                first_paragraph.paragraph_format.space_after = \
                    Pixel(space_after, self.dpi.vertical_resolution)
            else:
                # If the space after not greater than 0
                # using first paragraph for generating text
                p = first_paragraph
            first_paragraph_available = False
        else:
            p = document.add_paragraph()

        if par.is_list and len(par.lines) > 0 and len(par.lines[0].words) > 1:
            self.is_list_para = update_bullets_and_numbering(
                self,
                document,
                p,
                par,
                carea_x,
                Pixel,
                is_near_image,
                box_container
            )
        else:
            self.is_list_para = False

        self.generate_paragraph(p,
                                par,
                                self.is_list_para,
                                box_container,
                                block)
        # Analysis text angle
        text_angle_histogram = \
            Common.get_text_angle_histogram(par, text_angle_histogram)

        self._set_space_after(page, block, box_container,
                              p, par, paragraph_orders, i)
        return p, first_paragraph_available, text_angle_histogram

    def generate_carea(
            self,
            document,
            page,
            block,
            first_paragraph_available,
            box_container):
        """
        Generate a content area into MS Word Document representation

        Parameters
        ----------
        document
        page
        block
        first_paragraph_available
        box_container

        Returns
        -------

        """
        angle = ''
        carea_xml = ''
        # Firstly, store position of each paragraph of content area
        paragraph_orders = Common.get_paragraph_array(
            block,
            self.text_direction
        )
        i = 0
        text_angle_histogram = {}
        for para_obj in paragraph_orders:
            for par in block.paragraphs:
                if par.id == para_obj.obj_id:
                    p, first_paragraph_available, text_angle_histogram = \
                        self.generate_paragraph_inside_carea(
                            document,
                            page,
                            block,
                            first_paragraph_available,
                            box_container,
                            par,
                            text_angle_histogram,
                            paragraph_orders,
                            i)
                    carea_xml += p._element.xml
                    if not box_container.is_whole_page_box:
                        # Clear data to use only textbox
                        Common.delete_paragraph(p)
                    i += 1
                    break

        angle = Common.text_angle_in_block(angle, text_angle_histogram)
        return carea_xml, angle

    def generate_tables(self, document, page):
        for table in page.tables:
            # Assign columns' width and height
            column_width, column_height = \
                PHOcrCommon.get_columns_size(table)
            # Generate table based on width, height, and
            # words inside of it
            self.generate_table(document, page, table,
                                column_width, column_height)

    def generate_complicated_document(self, document, page, page_box):
        """
        Generate a word document representation in case complicated layout

        Parameters
        ----------
        document : Document
            Instance of Document object in docx library
        page : Object
            XML object information of page
        page_box : OCRBox
            Position of page by bounding box
        Returns
        -------

        """
        # We need to generate table first to avoid conflicting between
        # table and photo, from that, lead to break page
        self.generate_tables(document, page)
        for photo in page.photos:
            photo_wrap = 'wrapNone'
            self.generate_photo(document, photo, photo_wrap)
        for carea in page.careas:
            carea_x, carea_y, carea_w, carea_h = \
                text_box.calculate_text_box_size(carea)
            carea_container_box = OCRBox(
                self.dpi,
                carea_x, carea_y,
                carea_w, carea_h,
                Pixel
            )
            box_container = BoxContainer(
                box=carea_container_box,
                parent_box=page_box
            )
            paragraphs, angle = self.generate_carea(
                document,
                page,
                carea,
                False,
                box_container=box_container
            )
            textbox = TextBoxAttribute(
                self.dpi,
                carea_x,
                carea_y,
                carea_w,
                carea_h,
                paragraphs,
                angle
            )
            Common.add_textbox(document, textbox)

    def generate_carea_by_textbox(self, document, page, element, page_box):
        carea_x, carea_y, carea_w, carea_h = \
            text_box.calculate_text_box_size(
                element
            )
        carea_container_box = OCRBox(
            self.dpi,
            carea_x, carea_y,
            carea_w, carea_h,
            Pixel
        )
        box_container = BoxContainer(
            box=carea_container_box,
            parent_box=page_box
        )
        paragraphs, angle = self.generate_carea(
            document,
            page,
            element,
            False,
            box_container
        )
        textbox = TextBoxAttribute(
            self.dpi,
            carea_x,
            carea_y,
            carea_w,
            carea_h,
            paragraphs,
            angle
        )
        Common.add_textbox(document, textbox)

    def generate_normal_carea(
            self, document, page, element,
            page_box, first_para_available, first_textbox):
        carea_container_box = OCRBox(
            self.dpi,
            self.margin_left,
            self.margin_top,
            (self.width
             - self.margin_left
             - self.margin_right),
            (self.height
             - self.margin_top
             - self.margin_bottom),
            Pixel
        )
        box_container = BoxContainer(
            box=carea_container_box,
            parent_box=page_box,
            is_whole_page_box=True
        )
        self.generate_carea(document,
                            page,
                            element,
                            first_para_available,
                            box_container)
        first_para_available = False

        if first_textbox:
            element_y = int(
                self.transformer.get_y(element)
            )
            if self.transformer.text_line_order_ == TextlineOrder.TEXTLINE_ORDER_TOP_TO_BOTTOM:
                resolution = self.dpi.vertical_resolution
            else:
                resolution = self.dpi.horizontal_resolution
            element_y = Pixel(
                element_y,
                resolution
            ).inches
            spacing_before_inches = \
                element_y / self.get_top_margin_inches()
            first_textbox = False
        else:
            spacing_before_inches = 0
        if len(document.paragraphs) > 0:
            paragraphs_ = document.paragraphs[-1]
            paragraphs_.paragraph_format.space_before = \
                Inches(spacing_before_inches)
        return first_para_available, first_textbox

    def generate_normally_document(self, document, page, page_box):
        """
        Generate a word document representation using textbox due to complicated

        Parameters
        ----------
        document : Document
            Instance of Document object
        page : Object
            XML object of page
        page_box : OCRBox
            Position by bounding box of page
        Returns
        -------

        """
        # We need to generate table first to avoid confliction between
        # table and photo, from that, lead to break page
        self.generate_tables(document, page)

        # Always add first paragraph for purpose
        # 1. Anchor for photo
        # 2. Text paragraph will add to here
        if len(document.paragraphs) == 0:
            document.add_paragraph()
        for photo in page.photos:
            if photo.inside_table:
                photo_wrap = 'wrapNone'
            else:
                photo_wrap = Common.get_photo_wrap(photo, page)
            self.generate_photo(document, photo, photo_wrap)
        first_para_available = True
        first_textbox = False
        for i in range(len(self.area_orders)):
            for carea in page.careas:
                area_obj = self.area_orders[i]
                if carea.id == area_obj.obj_id:
                    if carea.is_noise or \
                            ratio_overlap_photo(
                                carea,
                                self.all_box_photo) >= 1:
                        self.generate_carea_by_textbox(
                            document, page,
                            carea, page_box)
                    else:
                        first_para_available, first_textbox = \
                            self.generate_normal_carea(
                                document,
                                page,
                                carea,
                                page_box,
                                first_para_available,
                                first_textbox
                            )
                    break

    def get_page_information(self, page):
        self.transformer.orientation = page.orientation
        self.transformer.writing_direction = page.writing_direction
        self.transformer.text_line_order = page.textline_order

        self.margin_left = page.left_margin
        self.margin_right = page.right_margin
        self.margin_top = page.top_margin
        self.margin_bottom = page.bottom_margin

        if page.baseline > 1:
            self.baseline = int(page.baseline)
        if page.x_height > 1:
            self.x_height = int(page.x_height)
        self.dpi = PHOcrCommon.get_dpi(page.dpi)
        self.paper_name = page.paper_name
        self.width = page.w
        self.height = page.h
        # Set page size to standard size
        self.is_using_standard_size = page.is_using_standard_size
        if page.is_using_standard_size:
            self.paper_width = page.paper_width
            self.paper_height = page.paper_height
            self.paper_code = page.paper_code
            self.paper_orientation = page.paper_orientation
        self.text_direction = TextDirection.LRTB

    def set_section_information(self, section, first_page):
        # Set up standard paper size and name
        # In case using the paper size of 1st page
        # we will use information of 1st paper size
        if self.paper_name == '1stPage':
            page_width = Inches(first_page.w / first_page.dpi.horizontal_resolution)
            page_height = Inches(first_page.h / first_page.dpi.vertical_resolution)
        else:
            page_width = Inches(self.width / self.dpi.horizontal_resolution)
            page_height = Inches(self.height / self.dpi.vertical_resolution)
        if self.is_using_standard_size:
            # Check if the code point is different with the default value
            if self.paper_code != 0:
                section.paper_size = self.paper_code
                page_width = Inches(self.paper_width)
                page_height = Inches(self.paper_height)
            if self.paper_orientation == "portrait":
                section.orientation = WD_ORIENT.PORTRAIT
            elif self.paper_orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = page_width
        section.page_height = page_height

        # If we found layout is Chinese or Japanese
        if self.transformer.writing_direction == 2 and \
                self.transformer.text_line_order == 1:
            if self.width > self.height:
                section.orientation = WD_ORIENT.LANDSCAPE
            else:
                section.orientation = WD_ORIENT.PORTRAIT
            section.text_direction = WD_TEXT_DIRECTION.TBRL

        # Measurement by emu
        section.top_margin = Pixel(self.margin_top, self.dpi.vertical_resolution)
        section.bottom_margin = Pixel(self.margin_bottom, self.dpi.vertical_resolution)
        section.left_margin = Pixel(self.margin_left, self.dpi.horizontal_resolution)
        section.right_margin = Pixel(self.margin_right, self.dpi.horizontal_resolution)

    def generate_document(self, office_obj):
        """
        Generate a word document representation

        Parameters
        ----------
        office_obj
            XML object which is parsed from .xml file

        Returns
        -------

        """
        # Create a MS Word document
        document = Document()
        # Information in "properties"
        """
        Expected result for all formats:
        - last modified: (blank)
        - created: (current date with format yyyy-mm-dd)
        - author: "None"
        - last modified by: "Not saved yet"
        """
        core_properties = document.core_properties
        from datetime import datetime
        core_properties.created = datetime.utcnow()
        core_properties.author = ''
        core_properties.comments = ''
        core_properties.last_modified_by = ''
        try:
            # Todo: set on for loop when we support multipage
            section = document.sections[0]
            for index, page in enumerate(office_obj.pages):
                self.get_page_information(page)
                self.set_section_information(section, office_obj.pages[0])
                # Store position of each area
                self.area_orders = Common.get_carea_array(page)

                # If we found layout is Chinese or Japanese
                if self.transformer.writing_direction == 2 and \
                        self.transformer.text_line_order == 1:
                    self.text_direction = TextDirection.TBRL
                    # Sort by x and reverse (DEC)
                    self.area_orders = sorted(
                        self.area_orders,
                        key=lambda e: e.x,
                        reverse=True
                    )

                self.all_box_photo = Common.get_photo_boxes(page)
                list_of_careas = Common.get_carea_array_base_on_photos(
                    page,
                    self.all_box_photo
                )
                self.is_complicated = \
                    Common.is_complicated_page_layout(
                        list_of_careas,
                        self.debug
                    )

                page_box = OCRBox.from_xml_element(page, self.dpi, Pixel)

                if self.text_direction == TextDirection.TBRL or \
                        not self.is_complicated:
                    self.generate_normally_document(document, page, page_box)
                else:
                    self.generate_complicated_document(document, page, page_box)
                Common.smooth_par_contain_background_indent(document)
                if index != (len(office_obj.pages) - 1):
                    # New page section break
                    document.add_section()
                    p = document.add_paragraph()
                    run = p.add_run()
                    run.add_lastRenderedPageBreak()

            # Save the file
            document.save(self.output_file)
        except Exception as ex:
            # Logging error
            self.error_logger.log(
                OCRLogLevel.LOG_LEVEL_ERROR,
                'Error occurred. Arguments {0}.'.format(ex.args)
            )
            raise

    def get_top_margin_inches(self):
        return Pixel(self.margin_top, self.dpi.vertical_resolution).inches
